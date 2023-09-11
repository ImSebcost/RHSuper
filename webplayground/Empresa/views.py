
from django.shortcuts import render, get_object_or_404, redirect
from .models import Empleado, Horario, PagosCesantias, RecursoHumano, Salario, Vacaciones
from .forms import EmpleadoForm, EmpleadoSearchForm, HorarioForm, PagosCesantiasForm, RecursoHumanoForm, SalarioForm, VacacionesForm
from docx import Document
from django.http import FileResponse
from django.db.models import Q
#Empleado
def empleado_list(request):
    form = EmpleadoSearchForm(request.GET)
    empleados = Empleado.objects.all()

    if form.is_valid():
        search_term = form.cleaned_data.get('search_term')
        if search_term:
            empleados = empleados.filter(
                Q(nombre__icontains=search_term) | Q(apellido__icontains=search_term)
            )

    return render(request, 'Empleado/empleado_list.html', {'empleados': empleados, 'form': form})

def empleado_create(request):
    if request.method == 'POST':
        form = EmpleadoForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('empleado_list')
    else:
        form = EmpleadoForm()
    return render(request, 'Empleado/empleado_create.html', {'form': form})

def empleado_edit(request, pk):
    empleado = get_object_or_404(Empleado, pk=pk)
    if request.method == 'POST':
        form = EmpleadoForm(request.POST, instance=empleado)
        if form.is_valid():
            form.save()
            return redirect('empleado_list')
    else:
        form = EmpleadoForm(instance=empleado)
    return render(request, 'Empleado/empleado_edit.html', {'form': form, 'empleado': empleado})

def empleado_delete(request, pk):
    empleado = get_object_or_404(Empleado, pk=pk)
    if request.method == 'POST':
        empleado.delete()
        return redirect('empleado_list')
    return render(request, 'Empleado/empleado_delete.html', {'empleado': empleado})

def horario_list(request):
    horarios = Horario.objects.all()
    return render(request, 'Horario/horario_list.html', {'horarios': horarios})

def horario_create(request):
    if request.method == 'POST':
        form = HorarioForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('horario_list')
    else:
        form = HorarioForm()
    return render(request, 'Horario/horario_create.html', {'form': form})

def horario_edit(request, pk):
    horario = get_object_or_404(Horario, pk=pk)
    if request.method == 'POST':
        form = HorarioForm(request.POST, instance=horario)
        if form.is_valid():
            form.save()
            return redirect('horario_list')
    else:
        form = HorarioForm(instance=horario)
    return render(request, 'Horario/horario_edit.html', {'form': form, 'horario': horario})



def salario_list(request):
    salarios = Salario.objects.all()
    return render(request, 'Salario/salario_list.html', {'salarios': salarios})

def salario_create(request):
    if request.method == 'POST':
        form = SalarioForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('salario_list')
    else:
        form = SalarioForm()
    return render(request, 'Salario/salario_create.html', {'form': form})

def salario_edit(request, pk):
    salario = Salario.objects.get(pk=pk)
    if request.method == 'POST':
        form = SalarioForm(request.POST, instance=salario)
        if form.is_valid():
            form.save()
            return redirect('salario_list')
    else:
        form = SalarioForm(instance=salario)
    return render(request, 'Salario/salario_list.html', {'form': form})

def salario_delete(request, pk):
    salario = Salario.objects.get(pk=pk)
    if request.method == 'POST':
        salario.delete()
        return redirect('salario_list')
    return render(request, 'Salario/salario_delete.html', {'salario': salario})

def pago_list(request):
    pagos = PagosCesantias.objects.all()
    return render(request, 'PagoCesantias/pago_list.html', {'pagos': pagos})

def pago_create(request):
    if request.method == 'POST':
        form = PagosCesantiasForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('pago_list')
    else:
        form = PagosCesantiasForm()
    return render(request, 'PagoCesantias/pago_create.html', {'form': form})

def pago_edit(request, pk):
    pago = get_object_or_404(PagosCesantias, pk=pk)
    if request.method == 'POST':
        form = PagosCesantiasForm(request.POST, instance=pago)
        if form.is_valid():
            form.save()
            return redirect('pago_list')
    else:
        form = PagosCesantiasForm(instance=pago)
    return render(request, 'PagoCesantias/pago_edit.html', {'form': form, 'pago': pago})

def pago_delete(request, pk):
    pago = get_object_or_404(PagosCesantias, pk=pk)
    if request.method == 'POST':
        pago.delete()
        return redirect('pago_list')
    return render(request, 'PagoCesantias/pago_delete.html', {'pago': pago})

def recurso_list(request):
    recursos = RecursoHumano.objects.all()
    return render(request, 'RecursoHumano/recurso_list.html', {'recursos': recursos})

def recurso_create(request):
    if request.method == 'POST':
        form = RecursoHumanoForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('recurso_list')
    else:
        form = RecursoHumanoForm()
    return render(request, 'RecursoHumano/recurso_create.html', {'form': form})

def recurso_edit(request, pk):
    recurso = get_object_or_404(RecursoHumano, pk=pk)
    if request.method == 'POST':
        form = RecursoHumanoForm(request.POST, instance=recurso)
        if form.is_valid():
            form.save()
            return redirect('recurso_list')
    else:
        form = RecursoHumanoForm(instance=recurso)
    return render(request, 'RecursoHumano/recurso_edit.html', {'form': form, 'recurso': recurso})

def recurso_delete(request, pk):
    recurso = get_object_or_404(RecursoHumano, pk=pk)
    if request.method == 'POST':
        recurso.delete()
        return redirect('recurso_list')
    return render(request, 'RecursoHumano/recurso_delete.html', {'recurso': recurso})

def vacaciones_list(request):
    vacaciones = Vacaciones.objects.all()
    return render(request, 'Vacaciones/vacaciones_list.html', {'vacaciones': vacaciones})

def vacaciones_create(request):
    if request.method == 'POST':
        form = VacacionesForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('vacaciones_list')
    else:
        form = VacacionesForm()
    return render(request, 'Vacaciones/vacaciones_create.html', {'form': form})

def vacaciones_edit(request, pk):
    vacacion = Vacaciones.objects.get(pk=pk)
    if request.method == 'POST':
        form = VacacionesForm(request.POST, instance=vacacion)
        if form.is_valid():
            form.save()
            return redirect('vacaciones_list')
    else:
        form = VacacionesForm(instance=vacacion)
    return render(request, 'Vacaciones/vacaciones_edit.html', {'form': form, 'vacacion': vacacion})

#Documento

def generate_employee_report(request, empleado_id):
    # Obtén el objeto Empleado
    empleado = Empleado.objects.get(pk=empleado_id)

    # Crear un documento de Word
    document = Document()

    # Agregar información del empleado
    document.add_heading(f'Informe de Empleado: {empleado.nombre} {empleado.apellido}', level=1)

    document.add_heading('Información del Empleado', level=2)
    document.add_paragraph(f'Nombre: {empleado.nombre}')
    document.add_paragraph(f'Apellido: {empleado.apellido}')
    document.add_paragraph(f'Fecha de Nacimiento: {empleado.fecha_nacimiento}')
    # Agregar más campos según tus modelos Empleado

    # Obtener información relacionada con el empleado
    horarios = Horario.objects.filter(empleado=empleado)
    vacaciones = Vacaciones.objects.filter(empleado=empleado)
    salarios = Salario.objects.filter(empleado=empleado)
    pagos_cesantias = PagosCesantias.objects.filter(empleado=empleado)
    recursos_humanos = RecursoHumano.objects.filter(departamento_perteneciente=empleado.departamento)

    # Agregar información de las tablas relacionadas
    document.add_heading('Horarios', level=2)
    for horario in horarios:
        document.add_paragraph(f'Día: {horario.dia_semana}')
        document.add_paragraph(f'Hora de Inicio: {horario.hora_inicio}')
        document.add_paragraph(f'Hora de Finalización: {horario.hora_finalizacion}')

    document.add_heading('Vacaciones', level=2)
    for vacacion in vacaciones:
        document.add_paragraph(f'Fecha de Inicio: {vacacion.fecha_inicio}')
        document.add_paragraph(f'Fecha de Finalización: {vacacion.fecha_finalizacion}')
        document.add_paragraph(f'Estado de Solicitud: {vacacion.estado_solicitud}')
        document.add_paragraph(f'Motivo: {vacacion.motivo}')

    document.add_heading('Salarios', level=2)
    for salario in salarios:
        document.add_paragraph(f'Fecha de Inicio del Período: {salario.fecha_inicio_periodo}')
        document.add_paragraph(f'Fecha de Fin del Período: {salario.fecha_fin_periodo}')
        document.add_paragraph(f'Salario Base: {salario.salario_base}')
        document.add_paragraph(f'Bonificaciones: {salario.bonificaciones}')
        document.add_paragraph(f'Deducciones: {salario.deducciones}')
        document.add_paragraph(f'Total a Pagar: {salario.total_pagar}')

    document.add_heading('Pagos de Cesantías', level=2)
    for pago in pagos_cesantias:
        document.add_paragraph(f'Fecha de Pago: {pago.fecha_pago}')
        document.add_paragraph(f'Valor Pagado: {pago.valor_pagado}')
        document.add_paragraph(f'Período Correspondiente: {pago.periodo_correspondiente}')

    document.add_heading('Recursos Humanos', level=2)
    for recurso in recursos_humanos:
        document.add_paragraph(f'Nombre del Recurso: {recurso.nombre_recurso}')
        document.add_paragraph(f'Descripción: {recurso.descripcion}')
        document.add_paragraph(f'Contacto: {recurso.contacto}')
        # Agregar más campos según tu modelo RecursoHumano

    # Guardar el documento en un archivo
    file_path = f'informe_empleado_{empleado_id}.docx'
    document.save(file_path)

    # Preparar la respuesta para descargar el archivo
    response = FileResponse(open(file_path, 'rb'))
    response['Content-Disposition'] = f'attachment; filename="{file_path}"'
    
    return response

